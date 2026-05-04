"""
clustering.py
-------------
K-Means clustering analysis for Shariah-compliance firm archetypes.

Workflow
--------
    select_optimal_k(df, cluster_vars)  ->  plots Elbow and Silhouette curves
    run_kmeans(df, cluster_vars, k)     ->  assigns cluster labels
    characterize_clusters(df_means)     ->  mean profiles per cluster
    generate_cluster_report(...)        ->  Word document with plots + tables
"""

import os
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.cluster import KMeans
from sklearn.preprocessing import StandardScaler
from sklearn.decomposition import PCA
from sklearn.metrics import silhouette_score

try:
    from kneed import KneeLocator
    _HAS_KNEED = True
except ImportError:
    _HAS_KNEED = False

DEFAULT_CLUSTER_VARS = ['ROA', 'DR', 'IAR', 'IncR', 'NLA', 'Log Firm Size']


# ---------------------------------------------------------------------------
# Optimal K selection
# ---------------------------------------------------------------------------

def select_optimal_k(
    df: pd.DataFrame,
    cluster_vars: list = DEFAULT_CLUSTER_VARS,
    k_max: int = 10,
) -> dict:
    """
    Determine the optimal number of clusters using the Elbow method (WCSS)
    and the Silhouette score.

    Args:
        df:            DataFrame with company-level observations.
        cluster_vars:  Variables to use for clustering.
        k_max:         Maximum K to evaluate.

    Returns:
        Dict with keys 'wcss', 'silhouette_scores', 'k_elbow', 'k_silhouette'.
    """
    df_means  = df.groupby('Company')[cluster_vars].mean().dropna()
    scaler    = StandardScaler()
    scaled    = scaler.fit_transform(df_means)

    wcss, s_scores = [], []

    for k in range(1, k_max + 1):
        km = KMeans(n_clusters=k, random_state=42, n_init=10)
        km.fit(scaled)
        wcss.append(km.inertia_)

    for k in range(2, k_max + 1):
        km     = KMeans(n_clusters=k, random_state=42, n_init=10)
        labels = km.fit_predict(scaled)
        s_scores.append(silhouette_score(scaled, labels))

    k_elbow       = None
    k_silhouette  = int(np.argmax(s_scores)) + 2  # s_scores starts at k=2

    if _HAS_KNEED:
        kl = KneeLocator(range(1, k_max + 1), wcss, S=1.0,
                         curve="convex", direction="decreasing")
        k_elbow = kl.elbow

    return {
        'wcss':              wcss,
        'silhouette_scores': s_scores,
        'k_elbow':           k_elbow,
        'k_silhouette':      k_silhouette,
        'df_means':          df_means,
        'scaled':            scaled,
        'scaler':            scaler,
    }


# ---------------------------------------------------------------------------
# K-Means application
# ---------------------------------------------------------------------------

def run_kmeans(
    df: pd.DataFrame,
    cluster_vars: list = DEFAULT_CLUSTER_VARS,
    k: int = 4,
) -> tuple[pd.DataFrame, StandardScaler]:
    """
    Fit K-Means with the specified K and return company-level cluster assignments.

    Args:
        df:            Panel DataFrame with 'Company' column.
        cluster_vars:  Variables used for clustering (company means).
        k:             Number of clusters.

    Returns:
        Tuple of (df_means with 'Cluster' column, fitted StandardScaler).
    """
    df_means = df.groupby('Company')[cluster_vars].mean().dropna()
    scaler   = StandardScaler()
    scaled   = scaler.fit_transform(df_means)

    km = KMeans(n_clusters=k, random_state=42, n_init=10)
    df_means['Cluster'] = km.fit_predict(scaled)

    # PCA for 2-D visualisation
    pca = PCA(n_components=2)
    pc  = pca.fit_transform(scaled)
    df_means['PC1'] = pc[:, 0]
    df_means['PC2'] = pc[:, 1]

    return df_means, scaler


# ---------------------------------------------------------------------------
# Archetype characterisation
# ---------------------------------------------------------------------------

def characterize_clusters(df_means: pd.DataFrame, cluster_vars: list) -> pd.DataFrame:
    """
    Compute mean values per cluster (sorted by mean ROA, descending).

    Args:
        df_means:     DataFrame with 'Cluster' column (from run_kmeans).
        cluster_vars: Variables to summarise.

    Returns:
        DataFrame with one row per cluster.
    """
    return (
        df_means.groupby('Cluster')[cluster_vars]
        .mean()
        .sort_values('ROA', ascending=False)
    )


# ---------------------------------------------------------------------------
# Full report generation
# ---------------------------------------------------------------------------

def generate_cluster_report(
    df: pd.DataFrame,
    cluster_vars: list = DEFAULT_CLUSTER_VARS,
    k: int = 4,
    report_file: str = "results/11_kmeans_clustering_report.docx",
) -> None:
    """
    Run K-Means, produce all visualisations, and compile a Word document.

    Args:
        df:            Panel DataFrame with 'Company' column.
        cluster_vars:  Variables used for clustering.
        k:             Number of clusters to use.
        report_file:   Destination path for the Word document.
    """
    import docx
    from docx.shared import Inches

    os.makedirs(os.path.dirname(report_file), exist_ok=True)
    tmp_dir = os.path.dirname(report_file)

    # --- Fit ---
    df_means, scaler = run_kmeans(df, cluster_vars, k)
    scaled = scaler.transform(df_means[cluster_vars])

    archetype_df    = characterize_clusters(df_means, cluster_vars)
    firm_counts     = df_means.groupby('Cluster').size().rename('Count')

    doc = docx.Document()
    doc.add_heading(f'K-Means Clustering Analysis (K = {k})', level=1)

    def _save_fig(fig, name):
        path = os.path.join(tmp_dir, f"_tmp_{name}.png")
        fig.savefig(path, bbox_inches='tight')
        plt.close(fig)
        return path

    def _embed(doc, fig, title, caption, w=6):
        doc.add_heading(title, level=3)
        path = _save_fig(fig, title.replace(" ", "_"))
        doc.add_picture(path, width=Inches(w))
        doc.add_paragraph(caption)
        os.remove(path)

    # --- PCA scatter ---
    fig_pca = plt.figure(figsize=(10, 8))
    sns.scatterplot(x='PC1', y='PC2', hue='Cluster', data=df_means,
                    palette='viridis', s=100, alpha=0.7, edgecolor='k')
    plt.title(f'Company Clusters (K={k}) — PCA Projection')
    plt.grid(True, linestyle='--', alpha=0.6)
    plt.tight_layout()
    _embed(doc, fig_pca, 'PCA Cluster Plot',
           f'Companies projected onto two principal components, coloured by K-Means cluster (K={k}).')

    # --- Archetype table ---
    doc.add_heading('Cluster Archetypes (Mean Values)', level=2)
    df_out = archetype_df.reset_index()
    t = doc.add_table(rows=df_out.shape[0] + 1, cols=df_out.shape[1])
    t.style = 'Table Grid'
    for j, col in enumerate(df_out.columns):
        t.cell(0, j).text = col
    for i, row in df_out.iterrows():
        for j, val in enumerate(row):
            t.cell(i + 1, j).text = f"{val:.4f}" if isinstance(val, float) else str(val)

    # --- Firm counts ---
    doc.add_heading('Firms per Cluster', level=2)
    t2 = doc.add_table(rows=len(firm_counts) + 1, cols=2)
    t2.style = 'Table Grid'
    t2.cell(0, 0).text = 'Cluster'
    t2.cell(0, 1).text = 'Count'
    for i, (cl, cnt) in enumerate(firm_counts.items()):
        t2.cell(i + 1, 0).text = str(cl)
        t2.cell(i + 1, 1).text = str(int(cnt))

    # --- Box plots ---
    df_unscaled = pd.DataFrame(
        scaler.inverse_transform(scaled),
        columns=cluster_vars,
        index=df_means.index,
    )
    df_unscaled['Cluster'] = df_means['Cluster'].values

    doc.add_heading('Variable Distributions by Cluster', level=2)
    for var in cluster_vars:
        fig_box = plt.figure(figsize=(10, 6))
        sns.boxplot(x='Cluster', y=var, data=df_unscaled,
                    hue='Cluster', palette='viridis', legend=False)
        plt.title(f'{var} by Cluster')
        plt.grid(axis='y', linestyle='--', alpha=0.7)
        plt.tight_layout()
        _embed(doc, fig_box, f'Box Plot: {var}',
               f'Distribution of {var} across K-Means clusters.')

    doc.save(report_file)
    print(f"Clustering report → {report_file}")

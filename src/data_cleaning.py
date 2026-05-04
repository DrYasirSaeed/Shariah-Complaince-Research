"""
data_cleaning.py
----------------
Combines the three period CSVs into a balanced panel, applies
winsorization, sector harmonization, and log-transforms Firm Size.

Pipeline
--------
    combine_and_compute_sales_growth(file1, file2, file3, out_file)
    winsorize_data(input_file, output_file, report_file)
    harmonize_sectors(input_file, output_file)
    log_transform_firm_size(input_file, output_file)
"""

import os
import numpy as np
import pandas as pd


# ---------------------------------------------------------------------------
# Step 4: Combine periods + Sales Growth
# ---------------------------------------------------------------------------

def combine_and_compute_sales_growth(
    file1: str,
    file2: str,
    file3: str,
    out_file: str,
) -> pd.DataFrame:
    """
    Vertically concatenate the three period CSVs, build a full Company×Year
    panel grid, forward/backward-fill sector metadata, and calculate
    year-over-year Sales Growth for each company.

    Args:
        file1:    Path to 01_computed_variables_2005_08.csv
        file2:    Path to 02_computed_variables_2009_13.csv
        file3:    Path to 03_computed_variables_2014_23.csv
        out_file: Output path for the combined CSV.

    Returns:
        Combined, sorted DataFrame.
    """
    def _clean(df: pd.DataFrame) -> pd.DataFrame:
        df.columns = df.columns.str.strip()
        df['Company'] = df['Company'].astype(str).str.strip()
        df['Year']    = df['Year'].astype(int)
        return df

    df1 = _clean(pd.read_csv(file1))
    df2 = _clean(pd.read_csv(file2))
    df3 = _clean(pd.read_csv(file3))

    df_combined_raw = pd.concat([df1, df2, df3], ignore_index=True)

    # Full Company × Year grid
    all_companies = df_combined_raw['Company'].unique()
    all_years     = range(2005, 2024)
    full_panel    = pd.DataFrame({
        'Company': np.repeat(all_companies, len(all_years)),
        'Year':    np.tile(all_years, len(all_companies)),
    })

    df_harmonized = pd.merge(full_panel, df_combined_raw,
                             on=['Company', 'Year'], how='left')

    # Fill Sector / Sub-Sector within each company
    for col in ('Sector', 'Sub-Sector'):
        df_harmonized[col] = (
            df_harmonized.groupby('Company')[col].ffill().bfill()
        )

    # Preserve original Excel company order
    orig_order = df1['Company'].drop_duplicates().tolist()
    for c in df_harmonized['Company'].drop_duplicates():
        if c not in orig_order:
            orig_order.append(c)

    df_harmonized['Company'] = pd.Categorical(
        df_harmonized['Company'], categories=orig_order, ordered=True
    )
    df_harmonized.sort_values(['Company', 'Year'], inplace=True)

    # Sales Growth
    df_harmonized['Sales'] = pd.to_numeric(df_harmonized['Sales'], errors='coerce')
    df_harmonized['Sales Growth'] = (
        df_harmonized
        .groupby('Company', observed=False)['Sales']
        .pct_change(fill_method=None)
    )
    df_harmonized['Sales Growth'] = df_harmonized['Sales Growth'].replace(
        [np.inf, -np.inf], np.nan
    )
    df_harmonized.loc[df_harmonized['Sales'] == 0, 'Sales Growth'] = np.nan

    # Column order
    lead_cols  = ['Sector', 'Sub-Sector', 'Company', 'Year']
    other_cols = [c for c in df_harmonized.columns if c not in lead_cols]
    df_harmonized = df_harmonized[lead_cols + other_cols]

    os.makedirs(os.path.dirname(out_file), exist_ok=True)
    df_harmonized.to_csv(out_file, index=False, encoding='utf-8-sig')
    print(f"Combined panel saved → {out_file}  ({len(df_harmonized):,} rows)")
    return df_harmonized


# ---------------------------------------------------------------------------
# Step 5: Winsorization
# ---------------------------------------------------------------------------

CONTINUOUS_COLS = [
    'ROA', 'DR', 'IR', 'IncR', 'NLA', 'IAR',
    'Firm Size', 'Tangibility', 'Sales', 'Sales Growth',
]


def winsorize_data(
    input_file: str,
    output_file: str,
    report_file: str | None = None,
) -> pd.DataFrame:
    """
    Winsorize continuous variables at the 1st and 99th percentiles.

    Args:
        input_file:  Path to the combined CSV (step 4 output).
        output_file: Path to save the winsorized CSV.
        report_file: Optional path to save a Word document summary.

    Returns:
        Winsorized DataFrame.
    """
    df = pd.read_csv(input_file)
    report_rows = []

    for col in CONTINUOUS_COLS:
        if col not in df.columns:
            continue
        df[col] = pd.to_numeric(df[col], errors='coerce')
        orig_min = df[col].min()
        orig_max = df[col].max()
        lo = df[col].quantile(0.01)
        hi = df[col].quantile(0.99)
        n_modified = int(((df[col] < lo) | (df[col] > hi)).sum())
        df[col] = df[col].clip(lower=lo, upper=hi)
        report_rows.append({
            'Variable':          col,
            'Original Min':      f"{orig_min:.4f}",
            'Original Max':      f"{orig_max:.4f}",
            '1st Pct (New Min)': f"{lo:.4f}",
            '99th Pct (New Max)':f"{hi:.4f}",
            'Values Modified':   str(n_modified),
        })

    os.makedirs(os.path.dirname(output_file), exist_ok=True)
    df.to_csv(output_file, index=False)
    print(f"Winsorized data saved → {output_file}")

    if report_file:
        _write_winsorization_report(report_rows, report_file)

    return df


def _write_winsorization_report(rows: list, report_file: str) -> None:
    import docx
    os.makedirs(os.path.dirname(report_file), exist_ok=True)
    doc = docx.Document()
    doc.add_heading('Winsorization Report', level=1)
    doc.add_paragraph(
        "Continuous variables winsorized at the 1st and 99th percentiles "
        "to mitigate extreme outlier influence while retaining all observations."
    )
    t = doc.add_table(rows=1, cols=6)
    t.style = 'Table Grid'
    hdrs = ['Variable', 'Original Min', 'Original Max',
            'New Min (1st Pct)', 'New Max (99th Pct)', 'Obs Modified']
    for i, h in enumerate(hdrs):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        cells[0].text = row['Variable']
        cells[1].text = row['Original Min']
        cells[2].text = row['Original Max']
        cells[3].text = row['1st Pct (New Min)']
        cells[4].text = row['99th Pct (New Max)']
        cells[5].text = row['Values Modified']
    doc.save(report_file)
    print(f"Winsorization report → {report_file}")


# ---------------------------------------------------------------------------
# Step 6: Sector harmonization and data cleaning
# ---------------------------------------------------------------------------

def harmonize_sectors(input_file: str, output_file: str) -> pd.DataFrame:
    """
    Apply targeted sector corrections and remove erroneous rows.

    Corrections applied:
      - Rename 'Food' → 'Food Products' across Sector and Sub-Sector.
      - Reclassify 'Goodluck Industries Ltd.' (years 2005–2013) to 'Food Products'.
      - Remove rows where Company == '727 - Food Products' (aggregate entry).

    Args:
        input_file:  Path to winsorized CSV (step 5 output).
        output_file: Path to save the cleaned CSV.

    Returns:
        Cleaned DataFrame.
    """
    df = pd.read_csv(input_file)

    # Rename 'Food' → 'Food Products'
    df['Sector']     = df['Sector'].replace('Food', 'Food Products')
    df['Sub-Sector'] = df['Sub-Sector'].replace('Food', 'Food Products')

    # Reclassify Goodluck Industries for 2005–2013
    goodluck = '460001 - Goodluck Industries Ltd.'
    mask = (df['Company'] == goodluck) & df['Year'].between(2005, 2013)
    df.loc[mask, ['Sector', 'Sub-Sector']] = 'Food Products'

    # Remove aggregate row
    df = df[df['Company'] != '727 - Food Products']

    os.makedirs(os.path.dirname(output_file), exist_ok=True)
    df.to_csv(output_file, index=False)
    print(f"Cleaned data saved → {output_file}  ({len(df):,} rows)")
    return df


# ---------------------------------------------------------------------------
# Step 7: Log-transform Firm Size
# ---------------------------------------------------------------------------

def log_transform_firm_size(input_file: str, output_file: str) -> pd.DataFrame:
    """
    Replace 'Firm Size' (total assets) with its natural log.

    Non-positive values are converted to NaN before the transformation.
    The original 'Firm Size' column is dropped.

    Args:
        input_file:  Path to the cleaned CSV (step 6 output).
        output_file: Path to save the transformed CSV.

    Returns:
        DataFrame with 'Log Firm Size' column.
    """
    df = pd.read_csv(input_file)
    df['Log Firm Size'] = df['Firm Size'].apply(
        lambda x: np.log(x) if pd.notna(x) and x > 0 else np.nan
    )
    df = df.drop(columns=['Firm Size'])

    os.makedirs(os.path.dirname(output_file), exist_ok=True)
    df.to_csv(output_file, index=False)
    print(f"Log-transformed data saved → {output_file}")
    return df

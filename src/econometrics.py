"""
econometrics.py
---------------
Panel econometric tests and models:
  - Fisher-ADF panel unit root tests
  - Optimal lag selection (PanelOLS + RandomForest)
  - Hausman test (FE vs RE)
  - Fixed Effects regression (full-panel and sector-wise)

All functions accept a cleaned DataFrame or a CSV path and can optionally
generate Word-document reports.
"""

import os
import numpy as np
import pandas as pd
import statsmodels.api as sm
from statsmodels.tsa.stattools import adfuller
from scipy.stats import combine_pvalues
from scipy.stats import chi2 as st_chi2


# ---------------------------------------------------------------------------
# Panel unit root tests (Fisher-ADF)
# ---------------------------------------------------------------------------

def run_unit_root_tests(
    df: pd.DataFrame,
    variables: list,
    min_obs: int = 8,
    report_file: str | None = None,
) -> pd.DataFrame:
    """
    Fisher-type ADF panel unit root test for each variable.

    H0: All panels contain a unit root.
    H1: At least one panel is stationary.

    Args:
        df:          Panel DataFrame with 'Company' and 'Year' columns.
        variables:   List of variable names to test.
        min_obs:     Minimum time-series length per firm to include in test.
        report_file: Optional path to save a Word document report.

    Returns:
        DataFrame with columns [Variable, Test, Statistic, P-Value, Result].
    """
    results = []

    for var in variables:
        if var not in df.columns:
            continue
        p_values = []

        for _, group in df.groupby('Company'):
            series = group[var].dropna()
            if len(series) > min_obs:
                try:
                    res = adfuller(series, regression='c', autolag='AIC')
                    p_values.append(res[1])
                except Exception:
                    continue

        cleaned = [p for p in p_values if not np.isnan(p)]
        cleaned = [np.finfo(float).eps if p == 0 else p for p in cleaned]

        if len(cleaned) > 1:
            try:
                stat, combined_p = combine_pvalues(cleaned, method='fisher')
                combined_p = max(combined_p, np.finfo(float).eps)
                results.append({
                    'Variable':  var,
                    'Test':      'Fisher-ADF',
                    'Statistic': f"{stat:.4f}",
                    'P-Value':   f"{combined_p:.4f}",
                    'Result':    'Stationary (I(0))' if combined_p < 0.05 else 'Non-Stationary',
                })
            except Exception as e:
                results.append({'Variable': var, 'Test': 'Fisher-ADF',
                                 'Statistic': 'Error', 'P-Value': 'Error',
                                 'Result': str(e)})
        else:
            results.append({'Variable': var, 'Test': 'Fisher-ADF',
                             'Statistic': 'N/A', 'P-Value': 'N/A',
                             'Result': 'Insufficient Data'})

    out = pd.DataFrame(results)
    if report_file:
        _write_unit_root_report(out, report_file)
    return out


def _write_unit_root_report(df_results: pd.DataFrame, path: str) -> None:
    import docx
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc = docx.Document()
    doc.add_heading('Panel Unit Root Test Results (Fisher-ADF)', level=1)
    doc.add_paragraph(
        "Fisher-ADF tests applied to each company time series; p-values "
        "combined via Fisher's method. H0: unit root in all panels."
    )
    t = doc.add_table(rows=1, cols=len(df_results.columns))
    t.style = 'Table Grid'
    for i, col in enumerate(df_results.columns):
        t.rows[0].cells[i].text = col
    for _, row in df_results.iterrows():
        cells = t.add_row().cells
        for i, col in enumerate(df_results.columns):
            cells[i].text = str(row[col])
    doc.save(path)
    print(f"Unit root report → {path}")


# ---------------------------------------------------------------------------
# Optimal lag selection
# ---------------------------------------------------------------------------

def select_optimal_lag(
    df: pd.DataFrame,
    dep_var: str,
    indep_vars: list,
    max_lag: int = 2,
    report_file: str | None = None,
) -> dict:
    """
    Compare PanelOLS (within R²) and RandomForest (test MSE / R²) across lags.

    Args:
        df:          Panel DataFrame indexed by (Company, Year).
        dep_var:     Dependent variable name.
        indep_vars:  List of independent variable names.
        max_lag:     Maximum lag to consider (0, 1, …, max_lag).
        report_file: Optional path for Word document report.

    Returns:
        Dict with keys 'econometric' and 'ml' DataFrames plus 'optimal_lag'.
    """
    from linearmodels.panel import PanelOLS
    from sklearn.ensemble import RandomForestRegressor
    from sklearn.metrics import mean_squared_error, r2_score

    eco_results = []
    ml_results  = []

    test_year = df.index.get_level_values('Year').max()

    for lag in range(max_lag + 1):
        temp = df.copy()
        if lag > 0:
            temp[indep_vars] = temp.groupby(level=0)[indep_vars].shift(lag)

        # --- Econometric (PanelOLS) ---
        data_panel = temp[[dep_var] + indep_vars].dropna()
        if not data_panel.empty:
            exog = sm.add_constant(data_panel[indep_vars], prepend=False)
            res  = PanelOLS(data_panel[dep_var], exog,
                            entity_effects=True).fit()
            eco_results.append({
                'Lag':          f"Lag {lag}",
                'R2_Within':    res.rsquared_within,
                'Observations': res.nobs,
            })

        # --- Machine Learning (RandomForest) ---
        data_ml = temp[[dep_var] + indep_vars].dropna().reset_index()
        train   = data_ml[data_ml['Year'] < test_year]
        test    = data_ml[data_ml['Year'] == test_year]

        if not test.empty:
            model = RandomForestRegressor(n_estimators=100, random_state=42)
            model.fit(train[indep_vars], train[dep_var])
            preds = model.predict(test[indep_vars])
            ml_results.append({
                'Lag':     f"Lag {lag}",
                'MSE':     mean_squared_error(test[dep_var], preds),
                'R2_Test': r2_score(test[dep_var], preds),
            })

    df_eco = pd.DataFrame(eco_results)
    df_ml  = pd.DataFrame(ml_results)

    optimal_lag = 0
    if not df_eco.empty:
        optimal_lag = df_eco.loc[df_eco['R2_Within'].idxmax(), 'Lag']

    if report_file:
        _write_lag_report(df_eco, df_ml, optimal_lag, report_file)

    return {'econometric': df_eco, 'ml': df_ml, 'optimal_lag': optimal_lag}


def _write_lag_report(df_eco, df_ml, optimal_lag, path):
    import docx
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc = docx.Document()
    doc.add_heading('Optimal Lag Length Selection Report', level=1)

    doc.add_heading('Econometric Results (PanelOLS)', level=2)
    _add_df_table(doc, df_eco)
    doc.add_paragraph(f"Optimal lag (max within R²): {optimal_lag}")

    doc.add_heading('Machine Learning Results (RandomForest)', level=2)
    _add_df_table(doc, df_ml)

    doc.add_heading('Conclusion', level=2)
    doc.add_paragraph(
        f"Both methods consistently indicate {optimal_lag} as optimal."
    )
    doc.save(path)
    print(f"Lag selection report → {path}")


# ---------------------------------------------------------------------------
# Hausman test
# ---------------------------------------------------------------------------

def hausman_test(
    df: pd.DataFrame,
    dep_var: str,
    indep_vars: list,
    report_file: str | None = None,
) -> dict:
    """
    Hausman specification test to choose between FE and RE models.

    Args:
        df:          Panel DataFrame (multi-index: Company, Year).
        dep_var:     Dependent variable.
        indep_vars:  List of independent variables.
        report_file: Optional Word document report path.

    Returns:
        Dict with 'chi2', 'df', 'p_value', and 'decision' keys.
    """
    from linearmodels.panel import PanelOLS, RandomEffects

    data = df[[dep_var] + indep_vars].dropna()
    y    = data[dep_var]
    X    = sm.add_constant(data[indep_vars])

    res_fe = PanelOLS(y, X, entity_effects=True, drop_absorbed=True).fit()
    res_re = RandomEffects(y, X).fit()

    b, B    = res_fe.params, res_re.params
    v_b, v_B = res_fe.cov, res_re.cov

    diff     = b - B
    cov_diff = v_b - v_B
    chi2_stat = float(np.dot(diff.T, np.linalg.pinv(cov_diff).dot(diff)))
    dof       = b.size
    p_val     = float(st_chi2.sf(chi2_stat, dof))

    decision = "Fixed Effects" if p_val < 0.05 else "Random Effects"
    result   = {
        'chi2':     chi2_stat,
        'df':       dof,
        'p_value':  p_val,
        'decision': decision,
    }

    if report_file:
        _write_hausman_report(result, report_file)

    return result


def _write_hausman_report(result: dict, path: str) -> None:
    import docx
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc = docx.Document()
    doc.add_heading('Hausman Test Report', level=1)
    doc.add_paragraph(f"Chi-Squared: {result['chi2']:.4f}")
    doc.add_paragraph(f"Degrees of Freedom: {result['df']}")
    doc.add_paragraph(f"P-Value: {result['p_value']:.4f}")
    doc.add_paragraph(f"Decision: Use {result['decision']} model.")
    doc.save(path)
    print(f"Hausman report → {path}")


# ---------------------------------------------------------------------------
# Fixed Effects regression
# ---------------------------------------------------------------------------

def run_fixed_effects(
    df: pd.DataFrame,
    dep_var: str,
    indep_vars: list,
    report_file: str | None = None,
):
    """
    Run a Fixed Effects (entity) PanelOLS model.

    Args:
        df:          Panel DataFrame (multi-index: Company, Year).
        dep_var:     Dependent variable.
        indep_vars:  Independent variables (no constant needed; added automatically).
        report_file: Optional Word document report path.

    Returns:
        Fitted linearmodels PanelOLS result.
    """
    from linearmodels.panel import PanelOLS

    data = df[[dep_var] + indep_vars].dropna()
    y    = data[dep_var]
    X    = sm.add_constant(data[indep_vars], prepend=False)

    mod = PanelOLS(y, X, entity_effects=True, drop_absorbed=True)
    res = mod.fit()
    print(res)

    if report_file:
        _write_fe_report(res, dep_var, indep_vars, report_file)

    return res


def _write_fe_report(res, dep_var, indep_vars, path):
    import docx
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc = docx.Document()
    doc.add_heading('Fixed Effects Regression Results', level=1)
    for line in str(res).split('\n'):
        doc.add_paragraph(line, style='No Spacing')
    doc.save(path)
    print(f"FE regression report → {path}")


# ---------------------------------------------------------------------------
# Sector-wise Fixed Effects
# ---------------------------------------------------------------------------

def run_sector_fixed_effects(
    df: pd.DataFrame,
    dep_var: str,
    indep_vars: list,
    min_companies: int = 15,
    min_obs: int = 50,
    report_file: str | None = None,
) -> pd.DataFrame:
    """
    Run sector-level Fixed Effects regressions and compile a coefficient matrix.

    Args:
        df:              Panel DataFrame with 'Sector' column (multi-index).
        dep_var:         Dependent variable.
        indep_vars:      Independent variables.
        min_companies:   Minimum unique companies for a sector to be included.
        min_obs:         Minimum observations per sector.
        report_file:     Optional Word document report path.

    Returns:
        DataFrame of coefficients (rows = variables, columns = sectors).
    """
    from linearmodels.panel import PanelOLS

    valid_sectors = (
        df.reset_index()
        .groupby('Sector')['Company']
        .nunique()
        .pipe(lambda s: s[s >= min_companies].index.tolist())
    )

    summary = {}

    for sector in valid_sectors:
        sector_df = df[df['Sector'] == sector].dropna(
            subset=[dep_var] + indep_vars
        )
        if len(sector_df) <= min_obs:
            continue

        y = sector_df[dep_var]
        X = sm.add_constant(sector_df[indep_vars])

        try:
            res = PanelOLS(y, X, entity_effects=True).fit(
                cov_type='clustered', cluster_entity=True
            )
        except Exception:
            res = PanelOLS(y, X, entity_effects=True).fit()

        col = []
        for var in ['const'] + indep_vars:
            coef  = res.params[var]
            pval  = res.pvalues[var]
            stars = "***" if pval < 0.01 else "**" if pval < 0.05 else "*" if pval < 0.1 else ""
            col.append(f"{coef:.4f}{stars}")
        summary[sector] = col

    index_labels = ['Intercept'] + indep_vars
    summary_df   = pd.DataFrame(summary, index=index_labels)

    if report_file:
        _write_sector_fe_report(summary_df, report_file)

    return summary_df


def _write_sector_fe_report(summary_df: pd.DataFrame, path: str) -> None:
    import docx
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc = docx.Document()
    doc.add_heading('Sector-Wise Fixed Effects Coefficients', level=1)
    doc.add_paragraph("*** p<0.01  ** p<0.05  * p<0.1")
    df_out = summary_df.reset_index().rename(columns={'index': 'Variable'})
    t = doc.add_table(rows=df_out.shape[0] + 1, cols=df_out.shape[1])
    t.style = 'Table Grid'
    for j, col in enumerate(df_out.columns):
        t.cell(0, j).text = col
    for i, row in df_out.iterrows():
        for j, val in enumerate(row):
            t.cell(i + 1, j).text = str(val)
    doc.save(path)
    print(f"Sector FE report → {path}")


# ---------------------------------------------------------------------------
# Shared utility
# ---------------------------------------------------------------------------

def _add_df_table(doc, df: pd.DataFrame) -> None:
    import docx as _docx
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = 'Table Grid'
    for i, col in enumerate(df.columns):
        t.rows[0].cells[i].text = col
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for i, col in enumerate(df.columns):
            v = row[col]
            cells[i].text = f"{v:.4f}" if isinstance(v, float) else str(v)

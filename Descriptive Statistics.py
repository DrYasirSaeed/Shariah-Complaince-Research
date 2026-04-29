import pandas as pd
import docx
import numpy as np

# 1. Load the cleaned data
file_path = "final_unbalanced_panel_2026.csv"
df = pd.read_csv(file_path)

# 2. Define Variables
cols = ['ROA', 'DR', 'IR', 'IncR', 'NLA', 'IAR']
for c in cols:
    df[c] = pd.to_numeric(df[c], errors='coerce')

# 3. Create a custom grouping angle: Regulatory Period
df['Regulatory_Period'] = df['Year'].apply(lambda x: 'Pre-2014' if x < 2014 else 'Post-2014')

# 4. Helper function to format values to 4 decimal places
def fmt(val):
    if pd.isna(val): return "-"
    return f"{val:.4f}"

# 5. Initialize Word Document
doc = docx.Document()
doc.add_heading('Multi-Dimensional Descriptive Statistics', level=1)
doc.add_paragraph("This document provides a comprehensive breakdown of descriptive statistics across varying levels of granularity: Overall, Year-wise, Regulatory Period, Sector, Sub-Sector, and Company.")

# 6. Reusable Function to generate tables for different groupings
def add_grouped_mean_table(doc, df, groupby_col, cols, title, description):
    doc.add_heading(title, level=2)
    doc.add_paragraph(description)
    
    # Calculate group means and counts
    grouped = df.groupby(groupby_col)[cols].mean()
    counts = df.groupby(groupby_col).size()
    
    # Setup Word Table
    t = doc.add_table(rows=1, cols=len(cols) + 2)
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    hdr[0].text = groupby_col
    hdr[1].text = 'Obs (N)'
    for i, col in enumerate(cols):
        hdr[i+2].text = col
        
    # Populate rows
    for idx in grouped.index:
        # Skip the aggregate index row if present
        if str(idx) == 'All Sector': continue 
            
        row = t.add_row().cells
        row[0].text = str(idx)
        row[1].text = str(counts.loc[idx])
        for i, col in enumerate(cols):
            row[i+2].text = fmt(grouped.loc[idx, col])

# ---------------------------------------------------------
# ANGLE 1: Overall (Full descriptive stats)
# ---------------------------------------------------------
doc.add_heading('1. Overall Descriptive Statistics', level=2)
desc_overall = df[cols].agg(['count', 'mean', 'std', 'min', 'max']).T
t1 = doc.add_table(rows=1, cols=6)
t1.style = 'Table Grid'
h1 = t1.rows[0].cells
headers1 = ['Variable', 'Observations', 'Mean', 'Std. Dev.', 'Minimum', 'Maximum']
for i, h in enumerate(headers1): h1[i].text = h
for var in cols:
    row = t1.add_row().cells
    row[0].text = var
    row[1].text = str(int(desc_overall.loc[var, 'count']))
    row[2].text = fmt(desc_overall.loc[var, 'mean'])
    row[3].text = fmt(desc_overall.loc[var, 'std'])
    row[4].text = fmt(desc_overall.loc[var, 'min'])
    row[5].text = fmt(desc_overall.loc[var, 'max'])

# ---------------------------------------------------------
# DYNAMIC ANGLES (Using our helper function)
# ---------------------------------------------------------
# Angle 2: Year-wise
add_grouped_mean_table(doc, df, 'Year', cols, '2. Year-wise Distribution', "Annual means.")

# Angle 3: Regulatory Period
add_grouped_mean_table(doc, df, 'Regulatory_Period', cols, '3. Regulatory Period', "Pre vs Post 2014.")

# Angle 4: Sector
add_grouped_mean_table(doc, df, 'Sector', cols, '4. Sector Distribution', "Macro-industry means.")

# Angle 5: Sub-Sector
add_grouped_mean_table(doc, df, 'Sub-Sector', cols, '5. Sub-Sector Distribution', "Granular industry variations.")

# Angle 6: Company
add_grouped_mean_table(doc, df, 'Company', cols, '6. Company-wise Distribution', "Firm-level averages over the sample period.")

# Save the final file
doc.save("Multi_Dimensional_Descriptive_Stats.docx")
print("Document successfully generated.")

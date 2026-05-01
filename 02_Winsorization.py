!pip install python-docx
import pandas as pd
import docx
import os # Import os module for path checking
from google.colab import drive

# 1. Mount Google Drive to ensure file access
drive.mount('/content/drive')

# 2. Input and Output definitions
base_path = '/content/drive/MyDrive/Github-Cursor/Shariah Compliance Project/'
input_file = base_path + 'Extracted Data/04_computed_variables_2005_23_Combined.csv'

# Corrected output_file location to 'Extracted Data' folder
output_data_dir = base_path + "Extracted Data/"
output_file = output_data_dir + "05_Winsorized_Data.csv"

# Corrected report_file location to 'Results' folder with new name
report_data_dir = base_path + "Results/"
report_file = report_data_dir + "02_winsorization report.docx"

# Ensure directories exist
os.makedirs(output_data_dir, exist_ok=True)
os.makedirs(report_data_dir, exist_ok=True)

# 3. Load the dataset
# Add a check to ensure the input file exists
if not os.path.exists(input_file):
    raise FileNotFoundError(f"The input file was not found: {input_file}. Please ensure that the previous cell (Part 4) was run successfully to create this file.")
df = pd.read_csv(input_file)
continuous_cols = ['ROA', 'DR', 'IR', 'IncR', 'NLA', 'IAR', 'Firm Size', 'Tangibility', 'Sales', 'Sales Growth']

# Ensure numeric format
for col in continuous_cols:
    df[col] = pd.to_numeric(df[col], errors='coerce')

# 4. Track changes for the report
report_data = []

for col in continuous_cols:
    # Capture original extremes
    orig_min = df[col].min()
    orig_max = df[col].max()

    # Calculate boundaries
    lower_bound = df[col].quantile(0.01)
    upper_bound = df[col].quantile(0.99)

    # Count how many records will be modified
    lower_capped_count = (df[col] < lower_bound).sum()
    upper_capped_count = (df[col] > upper_bound).sum()
    total_modified = lower_capped_count + upper_capped_count

    # Apply Winsorization
    df[col] = df[col].clip(lower=lower_bound, upper=upper_bound)

    # Capture new extremes
    new_min = df[col].min()
    new_max = df[col].max()

    # Log data for the report
    report_data.append({
        'Variable': col,
        'Original Min': f"{orig_min:.4f}",
        'Original Max': f"{orig_max:.4f}",
        '1st Pct (New Min)': f"{new_min:.4f}",
        '99th Pct (New Max)': f"{new_max:.4f}",
        'Values Modified': str(total_modified)
    })

# Save clean data
df.to_csv(output_file, index=False)

# 5. Generate the Word Document Report
doc = docx.Document()
doc.add_heading('Methodology: Data Treatment and Winsorization Report', level=1)

doc.add_paragraph(
    "To ensure the robustness of the econometric estimates and to mitigate the undue influence of extreme outliers, "
    "a systematic data treatment protocol was executed. Continuous financial variables exhibiting severe skewness—often "
    "resulting from near-zero denominators during periods of corporate distress—were Winsorized at the 1st and 99th percentiles."
)

doc.add_paragraph(
    "This procedure retains the underlying observations to prevent survivorship bias but caps extreme values at the specified "
    "percentile boundaries. Table 1 details the original boundaries, the newly established limits, and the exact number of "
    "firm-year observations modified per variable."
)

doc.add_heading('Table 1: Summary of Outlier Treatment (1% - 99% Winsorization)', level=2)

t = doc.add_table(rows=1, cols=6)
t.style = 'Table Grid'
hdr_cells = t.rows[0].cells
headers = ['Variable', 'Original Min', 'Original Max', 'New Min (1st Pct)', 'New Max (99th Pct)', 'Observations Modified']
for i, h in enumerate(headers):
    hdr_cells[i].text = h

for row_data in report_data:
    row_cells = t.add_row().cells
    row_cells[0].text = row_data['Variable']
    row_cells[1].text = row_data['Original Min']
    row_cells[2].text = row_data['Original Max']
    row_cells[3].text = row_data['1st Pct (New Min)']
    row_cells[4].text = row_data['99th Pct (New Max)']
    row_cells[5].text = row_data['Values Modified']

doc.save(report_file)
print(f"Data successfully cleaned and report generated: {report_file}")

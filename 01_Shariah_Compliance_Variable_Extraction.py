import pandas as pd
import numpy as np
from google.colab import drive

# 1. Mount Google Drive
# This step connects your Google Colab environment to your Google Drive,
# allowing the script to access files stored there.
drive.mount('/content/drive')

# 2. File Paths
# Define the base path and the paths to the three CSV files generated in previous steps,
# as well as the output path for the final combined CSV.
base_path = '/content/drive/MyDrive/Github-Cursor/Shariah Compliance Project/Extracted Data/'

file1 = base_path + '01_computed_variables_2005_08.csv'
file2 = base_path + '02_computed_variables_2009_13.csv'
file3 = base_path + '03_computed_variables_2014_23.csv'
out_file = base_path + '04_computed_variables_2005_23_Combined.csv'

print("Loading the three CSV files...")
# Load the dataframes from the individual CSV files.
# A try-except block is used to gracefully handle cases where a file might be missing.
try:
    df1 = pd.read_csv(file1)
    df2 = pd.read_csv(file2)
    df3 = pd.read_csv(file3)
except FileNotFoundError as e:
    print(f"Error finding file: {e}")
    print("Please make sure all three scripts have been successfully run first.")
    raise

# --- Debugging: Print head of each DataFrame after loading ---
print("\nPreview of df1 (2005-08 data):")
display(df1.head())
print("\nPreview of df2 (2009-13 data):")
display(df2.head())
print("\nPreview of df3 (2014-23 data):")
display(df3.head())
# --- End Debugging ---

# 3. Combine all rows into one massive dataset
print("Concatenating files...")
# Use pd.concat to stack the three DataFrames vertically.
# 'ignore_index=True' resets the index for the combined DataFrame.
df_combined = pd.concat([df1, df2, df3], ignore_index=True)

# 4. Extract the exact original top-to-bottom sequence of companies from File 1
# This step ensures that the final sorted output maintains the company order as it appeared in the original Excel file.
# drop_duplicates() preserves the first time it sees each company, effectively keeping the original Excel sequence intact.
original_company_sequence = df1['Company'].drop_duplicates().tolist()

# (Safety Check): If any new companies appeared in 2009-13 or 2014-23 that weren't in 2005-08,
# append them to the end of our sequence list so they aren't lost in the sorting.
all_unique_companies = df_combined['Company'].drop_duplicates().tolist()
for company in all_unique_companies:
    if company not in original_company_sequence:
        original_company_sequence.append(company)

print("Sorting by Original Company Sequence, then chronologically by Year...")
# 5. Convert the 'Company' column into a Categorical data type with our custom sequence
# This is a key step for custom sorting. By making 'Company' a categorical type with a defined order,
# Pandas will sort based on this order rather than alphabetical order.
df_combined['Company'] = pd.Categorical(
    df_combined['Company'],
    categories=original_company_sequence,
    ordered=True
)

# 6. Sort! Because 'Company' is categorical, it will sort by your Excel order, then by Year
# The inplace=True argument modifies the DataFrame directly.
df_combined.sort_values(by=['Company', 'Year'], inplace=True)

# --- Debugging: Check year distribution in combined DataFrame ---
print("\nYear distribution in df_combined after sorting:")
display(df_combined['Year'].value_counts().sort_index())
# --- End Debugging ---

# Calculate Sales Growth
df_combined['Sales Growth'] = df_combined.groupby('Company', observed=False)['Sales'].pct_change(fill_method=None)
# Clean up infinite values which occur when previous year's sales were 0 or very small
df_combined['Sales Growth'] = df_combined['Sales Growth'].replace([np.inf, -np.inf], np.nan)
# If current year's sales are 0, set Sales Growth to NaN as per user's request (treating 0 sales as 'missing' growth).
df_combined.loc[df_combined['Sales'] == 0, 'Sales Growth'] = np.nan
# For the first year of each company, Sales Growth will be NaN, which is appropriate.

# 7. Export the final combined file
# Save the sorted, combined DataFrame to a new CSV file.
# 'index=False' prevents Pandas from writing the DataFrame index as a column.
# 'encoding='utf-8-sig'' handles special characters correctly.
df_combined.to_csv(out_file, index=False, encoding='utf-8-sig')

print(f"\nSuccess! Created Final Master File: {out_file}")
print(f"Total Rows: {len(df_combined)}")

# Display a quick preview of the top 20 rows to verify Company #1 has all its years stacked
print("\nPreview of the top 20 rows:")
display(df_combined.head(20))


import docx
import os
from google.colab import drive

# Mount Google Drive (if not already mounted)
drive.mount('/content/drive')

# Define report path
report_data_dir = '/content/drive/MyDrive/Github-Cursor/Shariah Compliance Project/Results/'
report_file = report_data_dir + '01_data_extraction_report.docx'

# Ensure the directory exists
os.makedirs(report_data_dir, exist_ok=True)

doc = docx.Document()
doc.add_heading('Data Extraction and Calculation Methodology Report', level=1)

doc.add_paragraph(
    "This report details the methodology employed for extracting financial data from a multi-sheet Excel file "
    "(2005-23.xlsx), calculating Shariah-compliance variables, and combining the data into a unified dataset."
)

doc.add_heading('1. Data Extraction from Excel Sheets', level=2)
doc.add_paragraph(
    "Financial data was extracted from three distinct sheets within the Excel file: '2005-08', '2009-13', and '2014-23'. "
    "Each sheet represented a specific time period and exhibited variations in item naming conventions and column structures. "
    "The `process_shariah_data` function was designed to handle these variations, making the extraction modular and adaptable."
)
doc.add_paragraph(
    "The extraction process involved:"
)
doc.add_paragraph(
    "•   **Reading Excel Sheets:** `pd.read_excel(file_path, sheet_name=sheet_name, header=None)` was used to load data without assuming a header row."
    "•   **Data Cleaning:** Helper functions `normalize_item` and `parse_num` were applied to clean string values, remove extra whitespace, and convert financial figures (including percentages and numbers in parentheses for negatives) into a consistent float format."
    "•   **Item Name Matching:** The `get_item_year_value` function evolved to handle schema changes:"
    "    •   For '2005-08', it used exact item name matching."
    "    •   For '2009-13', it used lists of possible item names to account for minor variations."
    "    •   For '2014-23', it implemented a fuzzy matching logic that first tried exact matches, then partial case-insensitive matches to robustly identify financial items even with significant formatting differences (e.g., bracketed descriptions)."
)

doc.add_heading('2. Shariah-Compliance Variable Calculations', level=2)
doc.add_paragraph(
    "The following Shariah-compliance and control variables were calculated for each company-year observation based on the extracted financial items."
)

doc.add_heading('Table 1: Calculated Variables and Formulas', level=3)

table_data = [
    ['Variable', 'Description', 'Formula/Calculation'],
    ['ROA (Return on Assets)', 'Measures profitability relative to total assets.', 'Directly extracted or inferred from percentage values, then normalized to decimal. Original Excel definition: F7 as a % of Avg {Current year(A+B),previous year (A+B)}.'],
    ['DR (Debt Ratio)', 'Measures the proportion of assets financed by debt.', '(Non-Current Liabilities (D) + Short Term Borrowings (STB)) / Total Assets (TotalAB)'],
    ['IR (Investment Ratio)', 'Measures the proportion of assets held as short-term investments.', 'Short Term Investments (B4) / Total Assets (TotalAB)'],
    ['IncR (Income Ratio)', 'Measures the proportion of non-operating income relative to sales.', 'Other Income / (loss) (F5) / Sales (F1)'],
    ['IAR (Illiquid Asset Ratio)', 'Measures the proportion of illiquid assets relative to total assets.', '(Non-Current Assets (A) + Inventories (B2)) / Total Assets (TotalAB)'],
    ['NLA (Net Liquid Assets per Share)', 'Measures net liquid assets per share, indicating a company’s liquidity.', '(Total Assets (TotalAB) - Illiquid Assets (A + B2) - (Non-Current Liabilities (D) + Current Liabilities (E))) / (Issued, Subscribed & Paid up capital (C1) / 10.0)'],
    ['Firm Size', 'Proxy for company size.', 'Total Assets (TotalAB)'],
    ['Tangibility', 'Measures the proportion of assets that are tangible (non-current).', 'Non-Current Assets (A) / Total Assets (TotalAB)'],
    ['Sales', 'Raw sales figure.', 'Directly extracted Sales (F1)'],
    ['Non_Current_Assets', 'Raw non-current assets figure.', 'Directly extracted Non-Current Assets (A)'],
    ['Total_Assets', 'Raw total assets figure.', 'Directly extracted Total Assets (TotalAB)']
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
for i, item in enumerate(table_data[0]):
    hdr_cells[i].text = item

for row_items in table_data[1:]:
    row_cells = table.add_row().cells
    for i, item in enumerate(row_items):
        row_cells[i].text = item

doc.add_heading('3. Data Combination and Sales Growth Calculation', level=2)
doc.add_paragraph(
    "After processing each Excel sheet, the resulting CSV files (e.g., 01_computed_variables_2005_08.csv, etc.) "
    "were loaded and concatenated vertically using `pd.concat` to form a single master dataset. "
    "The original top-to-bottom sequence of companies from the first Excel sheet was preserved by converting the `Company` "
    "column to a Categorical data type with a predefined order, followed by sorting by `Company` and `Year`."
)

doc.add_paragraph(
    "**Sales Growth** was calculated as the percentage change in 'Sales' from the previous year for each company. "
    "Special handling was implemented to ensure accuracy for econometric analysis:"
)
doc.add_paragraph(
    "•   **Formula:** `df_combined.groupby('Company', observed=False)['Sales'].pct_change(fill_method=None)`"
    "•   **Infinite Values:** Infinite (`np.inf`, `-np.inf`) results, typically occurring when the previous year's sales were zero, were replaced with `np.nan`."
    "•   **Zero Sales:** If current year's sales were 0, 'Sales Growth' was explicitly set to `np.nan` to avoid misleading growth figures (e.g., 0/X or X/0) and accurately represent a lack of growth in such instances."
)

doc.save(report_file)
print(f"\nSuccess! Data extraction and calculation methodology report generated: {report_file}")

